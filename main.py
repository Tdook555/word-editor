from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
import uvicorn
import os
import json
import zipfile
from docx import Document
from typing import List

app = FastAPI()
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

def replace_in_paragraphs(paragraphs, replace_pairs):
    for para in paragraphs:
        for pair in replace_pairs:
            if not pair['old']:
                continue
            for run in para.runs:
                if pair['old'] in run.text:
                    run.text = run.text.replace(pair['old'], pair['new'])
            if pair['old'] in para.text:
                full_text = para.text
                new_text = full_text.replace(pair['old'], pair['new'])
                if para.runs:
                    para.runs[0].text = new_text
                    for run in para.runs[1:]:
                        run.text = ""

def count_in_paragraphs(paragraphs, replace_pairs):
    count = 0
    details = []
    for para in paragraphs:
        for pair in replace_pairs:
            if pair['old'] and pair['old'] in para.text:
                n = para.text.count(pair['old'])
                count += n
                details.append(f'"{pair["old"]}" → "{pair["new"]}" ({n} จุด)')
    return count, details

@app.get("/", response_class=HTMLResponse)
def index():
    with open("index.html", "r", encoding="utf-8") as f:
        return f.read()

@app.post("/preview")
async def preview_word(files: List[UploadFile] = File(...), pairs: str = Form(...)):
    replace_pairs = json.loads(pairs)
    results = []
    for file in files:
        input_path = f"{UPLOAD_DIR}/{file.filename}"
        content = await file.read()
        with open(input_path, "wb") as out_f:
            out_f.write(content)
        doc = Document(input_path)
        total = 0
        details = []

        c, d = count_in_paragraphs(doc.paragraphs, replace_pairs)
        total += c; details += d

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    c, d = count_in_paragraphs(cell.paragraphs, replace_pairs)
                    total += c; details += d

        for section in doc.sections:
            c, d = count_in_paragraphs(section.header.paragraphs, replace_pairs)
            total += c; details += d
            c, d = count_in_paragraphs(section.footer.paragraphs, replace_pairs)
            total += c; details += d

        results.append({"filename": file.filename, "total": total, "details": list(set(details))})

    return JSONResponse(results)

@app.post("/edit")
async def edit_word(files: List[UploadFile] = File(...), pairs: str = Form(...)):
    replace_pairs = json.loads(pairs)
    output_files = []
    for file in files:
        input_path = f"{UPLOAD_DIR}/{file.filename}"
        output_path = f"{UPLOAD_DIR}/edited_{file.filename}"
        content = await file.read()
        with open(input_path, "wb") as out_f:
            out_f.write(content)
        doc = Document(input_path)

        replace_in_paragraphs(doc.paragraphs, replace_pairs)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_paragraphs(cell.paragraphs, replace_pairs)

        for section in doc.sections:
            replace_in_paragraphs(section.header.paragraphs, replace_pairs)
            replace_in_paragraphs(section.footer.paragraphs, replace_pairs)

        doc.save(output_path)
        output_files.append(output_path)

    if len(output_files) == 1:
        return FileResponse(output_files[0], filename=f"edited_{files[0].filename}")

    zip_path = f"{UPLOAD_DIR}/edited_files.zip"
    with zipfile.ZipFile(zip_path, 'w') as zf:
        for f in output_files:
            zf.write(f, os.path.basename(f))

    return FileResponse(zip_path, filename="edited_files.zip")

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)